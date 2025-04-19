from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.oxml.text.font import CT_Fonts, CT_RPr
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.text.font import Font
from docx.document import Document as Document_Type

from bisect import bisect_left
from itertools import accumulate

from typing import NewType, Iterable

import re


Language = NewType('Language', str)
CHINESE = Language("c")
ENGLISH = Language("e")
ARABIC  = Language("a")
OTHERS  = Language("o")

Span = tuple[int, int]

def set_run_text(run_element: CT_R | Run, text: str) -> CT_R:
    """
    set the text of run_element to text

    Args:
        run_element (CT_R | Run): the run
        text (str): the text that u want the run element's text to change to

    Returns:
        run_element (CT_R): the same run_element but text changed
    """    
    run_element.text = text
    return run_element

def add_run_after_run(added_run: CT_R | Run | str, after_this_run: CT_R | Run, added_run_rPr: CT_RPr | Font | None = None) -> CT_R:
    """
    add the run "added_run" after the run "after_this_run".
    
    if "added_run" is a string, 
    the run added after "after_this_run" will be a run that 
    has the same font as "after_this_run" and text from the string "added_run".

    Args:
        added_run (CT_R | Run | str): added run or added text/string. if it is str and added_run_rPr type is not CT_RPr, font is inherited from after_this_run.
        after_this_run (CT_R | Run): the added_run will be added next to this run
        added_run_rPr (CT_RPr | Font | None): the wanted font for added_run. if None, font will be inherited from added_run, or from after_this_run only if added_run is str.
    
    Returns:
        added_run (CT_R): added run
    """
    if isinstance(added_run, str):
        added_run = set_run_text(after_this_run.__deepcopy__(""), added_run)
    elif isinstance(added_run, Run):
        added_run = added_run._element
        
    if isinstance(after_this_run, Run):
        after_this_run = after_this_run._element
    
    if isinstance(added_run_rPr, CT_RPr):
        added_run.replace(added_run.get_or_add_rPr(), added_run_rPr)
    elif isinstance(added_run_rPr, Font):
        added_run.replace(added_run.get_or_add_rPr(), added_run_rPr._element.rPr)
    
    after_this_run.addnext(added_run.__deepcopy__(""))
    
    return added_run

def remove_run(run: CT_R | Run):
    """
    remove the run from its parent paragraph.

    Args:
        run (CT_R | Run): run to be removed
    """    

    if isinstance(run, Run):
        run = run._element
    run.getparent().remove(run)

def split_run_at_string_index(run: CT_R | Run, index: int):
    """
    split the "split_run" at a given index of the run's text, 
    such that the text of first  splited part of the run is original text[:index],
    and       the text of second splited part of the run is original text[index:].

    Args:
        split_run (CT_R | Run): the run to be split
        index (int): the index of the splited position in the text of "split_run"

    Raises:
        IndexError: index should be >0 and <len(split_run.text), otherwise raise IndexError.
    """    

    if isinstance(run, Run):
        run = run._element
    
    text = run.text
    if index <= 0 or index >= len(text): raise IndexError(f"can't cut run \"{text}\", which has length {len(text)}, at index {index}.")

    set_run_text(run, text[:index])
    add_run_after_run(text[index:], run)

def delete_paragraph(paragraph: CT_P | Paragraph):
    """
    remove paragraph from its parent.

    Args:
        paragraph (Paragraph | CT_P): paragraph to be removed
    """    

    if isinstance(paragraph, Paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    else:
        p = paragraph
        p.getparent().remove(p)
    
def get_font_name(object: CT_Fonts | CT_R | Font | Run, language: Language) -> str | None:
    """
    get font name from rFont.
    
    In a rFont object,
    it can obtain 4 kinds of fonts, each for 1 of the following 4 languages:
    english, chinese, arabics, other languages.
    
    to get english font name, input "e" (or ENGLISH if you import * from DocxTools) in language arg.
    likewise,
    for chinese font name, input "c" (or CHINESE);
    for arabic  font name, input "a" (or ARABIC);
    for others  font name, input "o" (or OTHERS).

    Args:
        rFont (CT_Fonts | CT_R | Font | Run): the object that u want font name from.
        language (str): the language that you want the font name from.

    Raises:
        ValueError: wrong language arg.

    Returns:
        str: the desired font name.
    """
    
    if isinstance(object, Font):
        rFont = object._element.rPr.rFonts
    elif isinstance(object, Run):
        object = object._element
        rFont = object.rPr.rFonts
    elif isinstance(object, CT_R):
        rFont = object.rPr.rFonts
    
    if rFont is None: return None
    
    if language == CHINESE:
        return rFont.get(qn('w:eastAsia'))
    elif language == ENGLISH:
        return rFont.get(qn('w:ascii'))
    elif language == ARABIC:
        return rFont.get(qn('w:cs'))
    elif language == OTHERS:
        return rFont.get(qn('w:hAnsi'))
    else:
        raise ValueError("language must be c, e, a, or o.\n(stands for chinese, english, arabic or others respectively)")

def set_font_name(object: CT_Fonts | CT_R | Font | Run, new_font_name: str | dict[str], language: Language = ENGLISH) -> CT_Fonts:
    """
    set font name of object.
    
    In a rFont object,
    it can obtain 4 kinds of fonts, each for 1 of the following 4 languages:
    english, chinese, arabics, other languages.
    
    to set english font name, input "e" (or ENGLISH if you import * from DocxTools) in language arg.
    likewise,
    for chinese font name, input "c" (or CHINESE);
    for arabic  font name, input "a" (or ARABIC);
    for others  font name, input "o" (or OTHERS).

    Args:
        object (CT_Fonts | CT_R | Font | Run): the rFont object whose font name u want to set.
        new_font_name (str | dict[str]): if str, set font name to new_font_name;
                                         if dict[name], set font name to new_font_name[old_font_name].
        language (str): the language whose font name u want to set.

    Raises:
        ValueError: wrong language arg.

    Returns:
        CT_Fonts: the rFont object after its font name set.
    """  
    if isinstance(new_font_name, dict):
        replace_font_dict = new_font_name
        old_font_name = get_font_name(object, language)
        if old_font_name is None: return None
        if old_font_name in replace_font_dict.keys():
            new_font_name = replace_font_dict[old_font_name]
            return set_font_name(object, new_font_name, language)
        return None
    elif isinstance(new_font_name, str):
        if isinstance(object, Font):
            rFont = object._element.rPr.get_or_add_rFonts()
        elif isinstance(object, CT_R):
            rFont = object.rPr.get_or_add_rFonts()
        elif isinstance(object, Run):
            rFont = object._element.rPr.get_or_add_rFonts()
            
        if language == CHINESE:
            return rFont.set(qn('w:eastAsia'), new_font_name)
        elif language == ENGLISH:
            return rFont.set(qn('w:ascii'), new_font_name)
        elif language == ARABIC:
            return rFont.set(qn('w:cs'), new_font_name)
        elif language == OTHERS:
            return rFont.set(qn('w:hAnsi'), new_font_name)
        else:
            raise ValueError("language must be c, e, a, or o.\n(stands for chinese, english, arabic or others respectively)")

def find_reference_in_repl(repl: str) -> list[tuple[Span, str, re.Match[str]]]:
    pattern = re.compile(r"(?P<whole>\\(?P<nameORnumber>(?:\d+)|(?:g<\w+>)))")
    matches = list(pattern.finditer(repl))
    return [(m.span(), m.groupdict()['nameORnumber'], m) for m in matches]

def isolate_para_runs_by_span(paragraph: Paragraph | CT_P, 
                           span: Iterable[int]) -> tuple[int, int]:
    """
    split some of the given paragraph's runs, such that there exist a slice of paragraph.runs
    whose text contain and only contain the text of paragraph.text[span[0]:span[1]].
    
    return the tuple(starting index of the slice, ending index of the slice + 1).
    
    For example:
    
    I have a paragraph like this:
    string index                                                  0 1 2 3 4 5 6 7 8 9 10 
    paragraph text (separated into runs, indicated by |)          H e l|l o   W|o r l d
    run index                                                       0      1       2
    and the span is [4, 8], which mean I want to isolate paragraph text[4:8] i.e. "o Wo" in separate runs,
    
    then after this function is called, the given paragraph becomes like this:
    string index                                                  0 1 2 3 4 5 6 7 8 9 10
    paragraph text (separated into runs, indicated by |)          H e l|l|o   W|o|r l d
    run index                                                       0   1   2   3   4
    and (2, 4) will be returned, which indicates paragraph.runs[2:4] contains the isolated text "o Wo".

    Args:
        para_element (Paragraph | CT_P): paragraph to be split
        span (Iterable[int]): span of the para_element.text that is wanted in the section, i.e. you are trimming para_element.text[span[0]:span[1]].

    Returns:
        slice of paragraph.runs (tuple[int, int]): (starting index of the section, ending index of the section + 1)
    """    
    
    if isinstance(paragraph, Paragraph):
        paragraph = paragraph._element
    
    para_runs_start_index = [len(run_element.text) for run_element in paragraph.r_lst]
    para_runs_start_index = [0] + list(accumulate(para_runs_start_index))

    trim_start = span[0]                                              # index where the match start
    trim_end   = span[1]                                              # index where the match end
    trim_start_run = bisect_left(para_runs_start_index, trim_start)   # index in list para_runs_start_index if trim_start is to be inserted into the list
    trim_end_run   = bisect_left(para_runs_start_index, trim_end)     # index in list para_runs_start_index if trim_end   is to be inserted into the list

    if trim_end != para_runs_start_index[trim_end_run]:
        split_run_at_string_index(paragraph.r_lst[trim_end_run - 1],
                                            trim_end - para_runs_start_index[trim_end_run - 1])
    
    if trim_start != para_runs_start_index[trim_start_run]:
        split_run_at_string_index(paragraph.r_lst[trim_start_run - 1],
                                            trim_start - para_runs_start_index[trim_start_run - 1])
        trim_end_run += 1
    
    return trim_start_run, trim_end_run

def find(paragraph: CT_P, find: str) -> tuple[list[Span], list[list[Span]], list[re.Match[str]]]:
    """
    Find text in paragraph.
    
    First it finds the given texts in paragraph with regular expression.
    Then will isolate the found texts in the paragraph such that there exists runs that contain and only contain the found text (for details view MyDocxTools.isolate_para_runs_by_span).
    Finally return list of spans of the runs that contains the found text. length of the list equals to the number of occurance of the found text in the paragraph.

    Args:
        paragraph (CT_P): the paragraph to find text from
        find (str): the found text

    Returns:
        list[Span]: list of spans of the runs that contains the found text.
        
        list[list[Span]]]: In the second layer of list, which contain Spans, \
                           element at index 0 equals to the corresponding Span in the previous returned list[Span], \
                           which matches the whole regex expression, \
                           and the other elements at index > 0 matches corresponding captured group existing in the regex pattern. \
                           length of this list[0] == number of captured group existing in the regex pattern + 1.
                           
        list[re.Match[str]]: list of the match object
    """    
    para_text = paragraph.text
    matches = list(re.finditer(find, para_text))
    if matches == []: return [], [], []

    spans = [k.span() for k in matches]                  # e.g. if pattern = "a", text = "abc", then span = (0, 1)
    try:
        groups = [list(k.regs) for k in matches]
    except:
        groups = []
    
    for i, span in enumerate(spans):
        span = isolate_para_runs_by_span(paragraph, span)    
        
        # poss_groups_names = list(matches[i].groupdict.keys())
        # poss_groups_spans = [matches[i].span(n) for n in poss_groups_names]
        # groups_names = []
        
        prev_span_run_n = paragraph.r_lst.__len__()
        for ig, group in enumerate(groups[i][1:], start=1):
            # try:
            #     group_name = poss_groups_names[poss_groups_spans.index(group)]
            #     groups_names.append(group_name)
            # except:
            #     groups_names.append(None)
            groups[i][ig] = isolate_para_runs_by_span(paragraph, group)
        shift = paragraph.r_lst.__len__() - prev_span_run_n
        
        span = (span[0], span[1] + shift)
        spans[i] = span
        groups[i][0] = spans[i]
          
    return spans, groups, matches #, groups_names

def find_and_replace(paragraph_or_body: CT_P, finds: Iterable[str] | str, replaces: Iterable[str] | str, replaced_font: CT_RPr | Font | None = None) -> None:
    """
    You know... find and replace, from microsoft document.
    Just trying to replicate it with python-docx.
    
    Regular expression compatible.
    
    The font of the replaced text is inherited from the first run in the found text.

    Args:
        paragraph_or_body (CT_P | Paragraph | Document): the paragraph/document to undergo find and replace.
        finds (Iterable[str] | str):      find the strings from this list in the paragraph/document ...
        replaces (Iterable[str] | str):   and replace with the corresponding strings in this list.
    """
    if isinstance(paragraph_or_body, Document_Type):
        document = paragraph_or_body
        for paragraph in document._element.iterfind(".//" + qn("w:p")):
            find_and_replace(paragraph, finds, replaces)
    elif isinstance(paragraph_or_body, Paragraph):
        paragraph = paragraph_or_body
        find_and_replace(paragraph._element, finds, replaces)
    elif isinstance(paragraph_or_body, CT_P):
        paragraph = paragraph_or_body\
        
        if isinstance(finds, str):
            finds = [finds]
        if isinstance(replaces, str):
            replaces = [replaces]
        
        for find_text, replace_text in list(zip(finds, replaces)):
            
            repl_captured_groups = find_reference_in_repl(replace_text)
            
            repl_sections_names: list[None | str] = [None]
            repl_sections_matches: list[re.Match] = [None]
            repl_sections_starts = [0]
            # After the following for loop, repl_sections_names shall write down the names of the referenced groups that split the repl string into sections,
            # in the order of occurance in the repl, and shall write down None if no referenced group is in the section,
            # and repl_sections_starts shall drop down the start of each sections. Start of each section should also be the end of previous section.
            for repl_captured_group in repl_captured_groups:
                repl_group_span = repl_captured_group[0]
                repl_group_start = repl_group_span[0]
                repl_group_end = repl_group_span[1]
                repl_group_name = repl_captured_group[1]
                repl_group_match = repl_captured_group[2]
                
                if repl_group_start == repl_sections_starts[-1]:
                    repl_sections_names[-1] = repl_group_name
                else:
                    repl_sections_names.append(repl_group_name)
                    repl_sections_matches.append(repl_group_match)
                    repl_sections_starts.append(repl_group_start)
                    
                repl_sections_names.append(None)
                repl_sections_matches.append(None)
                repl_sections_starts.append(repl_group_end)
            repl_sections_ends = repl_sections_starts[1:] + [len(replace_text)]
            
            # Isolate all find_text and groups in paragraph
            spans, groupss, matches = find(paragraph, find_text)
            
            if matches:
                groupindex = matches[0].re.groupindex
            
            for i_spans in range(len(spans))[::-1]:
                match = matches[i_spans]
                groups = groupss[i_spans]
                run_span_start, run_span_end = spans[i_spans]
                
                # find the rPr of the first run in the span that is not in one of the groups
                # if no such run, give None
                fr = run_span_start
                first_normal_run_rPr = None
                for group in groups:
                    if group[0] == fr:
                        fr = group[1]
                    else:
                        first_normal_run_rPr = paragraph.r_lst[fr].rPr
                        break
                
                groups_rPr = [first_normal_run_rPr] + \
                             [paragraph.r_lst[run_group_start].rPr for run_group_start, _ in groups[1:]]
                
                # Remove unneeded runs
                run_to_remove = paragraph.r_lst[run_span_start: run_span_end]
                
                prev_run = paragraph.r_lst[run_span_start]
                for i, section_name in enumerate(repl_sections_names):
                    repl_section_string = replace_text[repl_sections_starts[i]: repl_sections_ends[i]]
                    repl_section_string = match.expand(repl_section_string)
                    
                    if section_name is None:
                        add_run_after_run(repl_section_string, prev_run, first_normal_run_rPr)
                    elif section_name.isnumeric():    
                        add_run_after_run(repl_section_string, prev_run, groups_rPr[int(section_name)])
                    else:
                        i_groups = groupindex[section_name]
                        group_rPr = groups_rPr[i_groups]
                        add_run_after_run(repl_section_string, prev_run, group_rPr)  # TODO change rPr to corresponding groups rPr
                
                for run in run_to_remove:
                    remove_run(run)
                    

if __name__ == "__main__":
    from docx import Document
    
    document = Document(r"../test.docx")
    find_and_replace(document, "(丹尼遜)(健腦操)(®)", r"\1\3\2")
    document.save(r"../output.docx")